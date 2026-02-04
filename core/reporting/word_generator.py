"""
Word Generator
==============
Generate professional Word documents using python-docx.

Features:
- RTL and Arabic support
- Tables with styling
- Headers and footers
- Page numbers
- Images
- Styles and formatting
- Templates support

Usage:
    from core.reporting import WordGenerator, WordConfig

    # Basic usage
    word = WordGenerator()
    word.add_heading("تقرير الموظفين")
    word.add_paragraph("هذا تقرير شامل عن الموظفين")
    word.add_table(data, headers=["الاسم", "الراتب"])
    word.save("report.docx")
"""

from dataclasses import dataclass
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from datetime import datetime
import io

from core.logging import app_logger


# Check python-docx availability
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False


@dataclass
class WordConfig:
    """Word generation configuration."""
    # Document info
    title: str = ""
    author: str = "INTEGRA"
    subject: str = ""

    # Page setup
    paper_size: str = "A4"
    orientation: str = "portrait"
    margins: Tuple[float, float, float, float] = (1.0, 1.0, 1.0, 1.0)  # inches

    # Styling
    rtl: bool = True
    font_family: str = "Cairo"
    font_size: int = 12
    heading_color: str = "#2563eb"
    primary_color: str = "#2563eb"

    # Header/Footer
    show_header: bool = True
    show_footer: bool = True
    show_page_numbers: bool = True
    header_text: str = ""
    footer_text: str = "INTEGRA - نظام الإدارة المتكامل"
    logo_path: Optional[str] = None


class WordGenerator:
    """
    Generate Word documents using python-docx.

    Supports RTL Arabic text, tables, images, and professional formatting.
    """

    def __init__(self, config: WordConfig = None):
        """
        Initialize Word generator.

        Args:
            config: Word configuration
        """
        if not WORD_AVAILABLE:
            raise ImportError(
                "python-docx not installed. Run: pip install python-docx"
            )

        self.config = config or WordConfig()
        self._document = Document()

        self._setup()
        app_logger.info("WordGenerator initialized")

    def _setup(self) -> None:
        """Setup document."""
        self._setup_page()
        self._setup_styles()
        self._setup_rtl()

    def _setup_page(self) -> None:
        """Setup page size and margins."""
        section = self._document.sections[0]

        # Orientation
        if self.config.orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            # Swap width and height for landscape
            section.page_width, section.page_height = section.page_height, section.page_width

        # Margins
        section.top_margin = Inches(self.config.margins[0])
        section.right_margin = Inches(self.config.margins[1])
        section.bottom_margin = Inches(self.config.margins[2])
        section.left_margin = Inches(self.config.margins[3])

    def _setup_styles(self) -> None:
        """Setup document styles."""
        styles = self._document.styles

        # Modify Normal style
        normal = styles['Normal']
        normal.font.name = self.config.font_family
        normal.font.size = Pt(self.config.font_size)

        # Set Arabic font
        normal._element.rPr.rFonts.set(qn('w:eastAsia'), self.config.font_family)

        # Create custom heading styles
        heading_sizes = {1: 24, 2: 20, 3: 16, 4: 14, 5: 12}

        for level, size in heading_sizes.items():
            try:
                heading_name = f'Heading {level}'
                if heading_name in styles:
                    heading = styles[heading_name]
                    heading.font.name = self.config.font_family
                    heading.font.size = Pt(size)
                    heading.font.bold = True
                    heading.font.color.rgb = RGBColor.from_string(
                        self.config.heading_color.lstrip('#')
                    )
            except Exception:
                pass

    def _setup_rtl(self) -> None:
        """Setup RTL support for Arabic."""
        if not self.config.rtl:
            return

        # Set document direction to RTL
        try:
            section = self._document.sections[0]

            # Get or create sectPr
            sectPr = section._sectPr

            # Set RTL direction using bidi
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')

            # Apply to section properties
            try:
                sectPr.append(bidi)
            except Exception:
                pass

        except Exception as e:
            app_logger.debug(f"RTL setup note: {e}")

    def _set_paragraph_rtl(self, paragraph) -> None:
        """Set RTL direction for a paragraph."""
        if not self.config.rtl:
            return

        try:
            pPr = paragraph._p.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            pPr.insert(0, bidi)

            # Set alignment to right
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'right')
            pPr.append(jc)
        except Exception:
            pass

    def _set_cell_rtl(self, cell) -> None:
        """Set RTL direction for a table cell."""
        if not self.config.rtl:
            return

        try:
            tcPr = cell._tc.get_or_add_tcPr()

            # Set text direction
            textDirection = OxmlElement('w:textDirection')
            textDirection.set(qn('w:val'), 'rlTb')
            tcPr.append(textDirection)

            # Set cell paragraphs to RTL
            for para in cell.paragraphs:
                self._set_paragraph_rtl(para)
        except Exception:
            pass

    def add_heading(
        self,
        text: str,
        level: int = 1,
        alignment: str = None
    ) -> 'WordGenerator':
        """
        Add heading to document.

        Args:
            text: Heading text
            level: Heading level (1-5)
            alignment: Override alignment (left, center, right)

        Returns:
            Self for chaining
        """
        heading = self._document.add_heading(text, level=level)

        # Set alignment
        align_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT
        }

        if alignment:
            heading.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.RIGHT)
        elif self.config.rtl:
            heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Set RTL
        self._set_paragraph_rtl(heading)

        return self

    def add_paragraph(
        self,
        text: str,
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
        font_size: int = None,
        alignment: str = None,
        color: str = None
    ) -> 'WordGenerator':
        """
        Add paragraph to document.

        Args:
            text: Paragraph text
            bold: Bold text
            italic: Italic text
            underline: Underlined text
            font_size: Font size in points
            alignment: Text alignment
            color: Text color (hex)

        Returns:
            Self for chaining
        """
        para = self._document.add_paragraph()

        # Set alignment
        align_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }

        if alignment:
            para.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.RIGHT)
        elif self.config.rtl:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add run with formatting
        run = para.add_run(text)
        run.font.name = self.config.font_family
        run.font.size = Pt(font_size or self.config.font_size)
        run.bold = bold
        run.italic = italic
        run.underline = underline

        if color:
            run.font.color.rgb = RGBColor.from_string(color.lstrip('#'))

        # Set RTL
        self._set_paragraph_rtl(para)

        return self

    def add_text(self, text: str, **kwargs) -> 'WordGenerator':
        """Alias for add_paragraph."""
        return self.add_paragraph(text, **kwargs)

    def add_table(
        self,
        data: List[Dict],
        headers: List[str] = None,
        style: str = "Table Grid"
    ) -> 'WordGenerator':
        """
        Add table to document.

        Args:
            data: List of row dictionaries
            headers: Column headers
            style: Table style name

        Returns:
            Self for chaining
        """
        if not data:
            return self

        # Get headers from first row if not provided
        if headers is None:
            headers = list(data[0].keys())

        # Create table
        num_cols = len(headers)
        table = self._document.add_table(rows=1, cols=num_cols)
        table.style = style
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # RTL table direction
        if self.config.rtl:
            try:
                tblPr = table._tbl.get_or_add_tblPr()
                bidiVisual = OxmlElement('w:bidiVisual')
                bidiVisual.set(qn('w:val'), '1')
                tblPr.append(bidiVisual)
            except Exception:
                pass

        # Add header row
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = str(header)
            self._set_cell_rtl(cell)

            # Style header
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.name = self.config.font_family

                # Background color
                try:
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), self.config.primary_color.lstrip('#'))
                    cell._tc.get_or_add_tcPr().append(shading)

                    # White text for header
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                except Exception:
                    pass

        # Add data rows
        for row_idx, row_data in enumerate(data):
            row_cells = table.add_row().cells

            for i, header in enumerate(headers):
                cell = row_cells[i]
                value = row_data.get(header, "")
                cell.text = str(value) if value is not None else ""
                self._set_cell_rtl(cell)

                # Style data cell
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = self.config.font_family

                # Alternating row color
                if row_idx % 2 == 1:
                    try:
                        shading = OxmlElement('w:shd')
                        shading.set(qn('w:fill'), 'F5F5F5')
                        cell._tc.get_or_add_tcPr().append(shading)
                    except Exception:
                        pass

        return self

    def add_image(
        self,
        image_path: str,
        width: float = None,
        height: float = None,
        alignment: str = "center"
    ) -> 'WordGenerator':
        """
        Add image to document.

        Args:
            image_path: Path to image file
            width: Image width in inches
            height: Image height in inches
            alignment: Image alignment

        Returns:
            Self for chaining
        """
        if not Path(image_path).exists():
            app_logger.warning(f"Image not found: {image_path}")
            return self

        try:
            para = self._document.add_paragraph()

            # Set alignment
            align_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT
            }
            para.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)

            run = para.add_run()

            if width:
                run.add_picture(image_path, width=Inches(width))
            elif height:
                run.add_picture(image_path, height=Inches(height))
            else:
                run.add_picture(image_path)

        except Exception as e:
            app_logger.error(f"Failed to add image: {e}")

        return self

    def add_line(self) -> 'WordGenerator':
        """Add horizontal line."""
        para = self._document.add_paragraph()
        para.add_run("─" * 60)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return self

    def add_spacer(self, lines: int = 1) -> 'WordGenerator':
        """Add empty lines."""
        for _ in range(lines):
            self._document.add_paragraph()
        return self

    def add_page_break(self) -> 'WordGenerator':
        """Add page break."""
        self._document.add_page_break()
        return self

    def add_header(self, text: str) -> 'WordGenerator':
        """
        Add document header.

        Args:
            text: Header text

        Returns:
            Self for chaining
        """
        section = self._document.sections[0]
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add line below header
        try:
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:color'), self.config.primary_color.lstrip('#'))
            pBdr.append(bottom)
            pPr.append(pBdr)
        except Exception:
            pass

        return self

    def add_footer(self, text: str, show_page_number: bool = True) -> 'WordGenerator':
        """
        Add document footer.

        Args:
            text: Footer text
            show_page_number: Include page numbers

        Returns:
            Self for chaining
        """
        section = self._document.sections[0]
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

        if show_page_number:
            para.text = f"{text}    |    "
            run = para.add_run()

            # Add page number field
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
        else:
            para.text = text

        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return self

    def save(self, output_path: str) -> bool:
        """
        Save document to file.

        Args:
            output_path: Output file path

        Returns:
            True if successful
        """
        try:
            # Ensure directory exists
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)

            # Set document properties
            core_props = self._document.core_properties
            core_props.title = self.config.title
            core_props.author = self.config.author
            core_props.subject = self.config.subject

            self._document.save(output_path)
            app_logger.info(f"Word document saved: {output_path}")
            return True

        except Exception as e:
            app_logger.error(f"Failed to save Word document: {e}", exc_info=True)
            return False

    def to_bytes(self) -> bytes:
        """
        Generate document as bytes.

        Returns:
            Document content as bytes
        """
        try:
            buffer = io.BytesIO()
            self._document.save(buffer)
            return buffer.getvalue()

        except Exception as e:
            app_logger.error(f"Failed to generate Word bytes: {e}")
            return b''

    def get_document(self) -> 'Document':
        """Get the underlying Document object."""
        return self._document


def create_word_report(
    data: List[Dict],
    output_path: str,
    title: str = "",
    headers: List[str] = None,
    **config_kwargs
) -> bool:
    """
    Quick function to create Word report.

    Args:
        data: Report data
        output_path: Output file path
        title: Report title
        headers: Table headers
        **config_kwargs: Additional WordConfig options

    Returns:
        True if successful
    """
    try:
        config = WordConfig(title=title, **config_kwargs)
        word = WordGenerator(config)

        # Add title
        if title:
            word.add_heading(title, level=1)
            word.add_spacer()

        # Add date
        word.add_paragraph(
            f"التاريخ: {datetime.now().strftime('%Y-%m-%d')}",
            alignment="left",
            font_size=10
        )
        word.add_line()
        word.add_spacer()

        # Add table
        if data:
            word.add_table(data, headers=headers)

        # Add footer
        if config.show_footer:
            word.add_spacer(2)
            word.add_paragraph(
                config.footer_text,
                alignment="center",
                font_size=10,
                color="#666666"
            )

        return word.save(output_path)

    except Exception as e:
        app_logger.error(f"Failed to create Word report: {e}")
        return False
