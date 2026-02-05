"""
Word Engine
===========
Word document reading, writing, and conversion.

Features:
- Read text and tables from .docx files
- Create new documents with headings, paragraphs, tables, images
- Save and convert to PDF
"""

from typing import List, Dict, Optional
from pathlib import Path

from core.logging import app_logger


class WordEngine:
    """Word document engine."""

    def __init__(self, file_path: str = None):
        """
        Initialize Word engine.

        Args:
            file_path: Path to existing .docx file, or None for new document
        """
        from docx import Document

        if file_path:
            self.doc = Document(file_path)
            self.file_path = file_path
        else:
            self.doc = Document()
            self.file_path = None

    # ═══════════════════════════════════════════════════════
    # Read Operations
    # ═══════════════════════════════════════════════════════

    def read_text(self) -> str:
        """Read all text from the document."""
        text = ""
        for para in self.doc.paragraphs:
            text += para.text + "\n"
        return text

    def read_tables(self) -> List[List[List[str]]]:
        """
        Read all tables from the document.

        Returns:
            List of tables, each table is a list of rows,
            each row is a list of cell texts.
        """
        tables = []
        for table in self.doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        return tables

    def get_headings(self) -> List[Dict]:
        """
        Get all headings from the document.

        Returns:
            List of dicts with heading text and level
        """
        headings = []
        for para in self.doc.paragraphs:
            if para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.replace('Heading ', ''))
                except ValueError:
                    level = 0
                headings.append({"text": para.text, "level": level})
        return headings

    def get_stats(self) -> Dict:
        """
        Get document statistics.

        Returns:
            Dict with word count, paragraph count, table count, etc.
        """
        text = self.read_text()
        words = text.split()

        return {
            "paragraphs": len(self.doc.paragraphs),
            "tables": len(self.doc.tables),
            "words": len(words),
            "characters": len(text),
            "sections": len(self.doc.sections),
        }

    # ═══════════════════════════════════════════════════════
    # Write Operations
    # ═══════════════════════════════════════════════════════

    def add_heading(self, text: str, level: int = 1):
        """
        Add a heading.

        Args:
            text: Heading text
            level: Heading level (0=Title, 1-9=Heading levels)
        """
        self.doc.add_heading(text, level=level)

    def add_paragraph(self, text: str, style: str = None,
                      alignment: str = "right"):
        """
        Add a paragraph.

        Args:
            text: Paragraph text
            style: Word paragraph style name
            alignment: 'left', 'center', 'right', 'justify'
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        para = self.doc.add_paragraph(text, style=style)

        alignment_map = {
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        if alignment in alignment_map:
            para.alignment = alignment_map[alignment]

    def add_table(self, data: List[List[str]],
                  headers: List[str] = None) -> None:
        """
        Add a table.

        Args:
            data: List of rows (list of cell values)
            headers: Optional header row
        """
        if headers:
            data = [headers] + data

        if not data or not data[0]:
            return

        table = self.doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Table Grid'

        for i, row_data in enumerate(data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    row.cells[j].text = str(cell_text)

    def add_image(self, image_path: str, width: float = None):
        """
        Add an image.

        Args:
            image_path: Path to the image file
            width: Width in inches (None for original size)
        """
        from docx.shared import Inches

        if width:
            self.doc.add_picture(image_path, width=Inches(width))
        else:
            self.doc.add_picture(image_path)

    def add_page_break(self):
        """Add a page break."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        para = self.doc.add_paragraph()
        run = para.add_run()
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        run._element.append(br)

    # ═══════════════════════════════════════════════════════
    # Save & Convert
    # ═══════════════════════════════════════════════════════

    def save(self, output_path: str = None) -> bool:
        """
        Save the document.

        Args:
            output_path: Output path (uses original path if None)

        Returns:
            True if successful
        """
        path = output_path or self.file_path
        if not path:
            return False

        try:
            self.doc.save(path)
            self.file_path = path
            return True
        except Exception as e:
            app_logger.error(f"Failed to save Word document: {e}")
            return False

    def to_pdf(self, output_path: str) -> bool:
        """
        Convert to PDF.

        Note: Requires docx2pdf or LibreOffice installed.

        Args:
            output_path: PDF output path

        Returns:
            True if successful
        """
        try:
            from docx2pdf import convert

            if not self.file_path:
                # Save to temp first
                import tempfile
                temp_path = tempfile.mktemp(suffix='.docx')
                self.save(temp_path)
                convert(temp_path, output_path)
                Path(temp_path).unlink(missing_ok=True)
            else:
                convert(self.file_path, output_path)

            return True
        except ImportError:
            app_logger.warning("docx2pdf not available, trying LibreOffice")
            return self._to_pdf_libreoffice(output_path)
        except Exception as e:
            app_logger.error(f"PDF conversion failed: {e}")
            return False

    def _to_pdf_libreoffice(self, output_path: str) -> bool:
        """Convert using LibreOffice as fallback."""
        import subprocess

        source = self.file_path
        if not source:
            return False

        try:
            output_dir = str(Path(output_path).parent)
            subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf',
                '--outdir', output_dir, source
            ], check=True, capture_output=True, timeout=60)

            # LibreOffice uses the same filename with .pdf extension
            lo_output = str(Path(output_dir) / (Path(source).stem + '.pdf'))
            if lo_output != output_path and Path(lo_output).exists():
                Path(lo_output).rename(output_path)

            return Path(output_path).exists()
        except Exception as e:
            app_logger.error(f"LibreOffice conversion failed: {e}")
            return False

    def to_text(self) -> str:
        """Convert document to plain text."""
        return self.read_text()
