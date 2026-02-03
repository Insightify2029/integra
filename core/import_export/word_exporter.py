"""
Word Exporter
=============
Export data to Word documents using python-docx.

Features:
- Create Word documents (.docx)
- Add headings, paragraphs, tables
- RTL support for Arabic
- Templates support
- Employee reports generation

Usage:
    from core.import_export import WordExporter

    # Simple document
    doc = WordExporter("report.docx")
    doc.add_heading("تقرير الموظفين", level=1)
    doc.add_paragraph("هذا تقرير شامل عن الموظفين")
    doc.add_table(data, headers=["الاسم", "القسم", "الراتب"])
    doc.save()

    # Employee report
    from core.import_export import create_employee_report
    create_employee_report(employee_data, "employee_report.docx")
"""

from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from core.logging import app_logger


class WordExporter:
    """Export data to Word documents."""

    def __init__(self, output_path: Optional[str] = None):
        """
        Initialize Word exporter.

        Args:
            output_path: Path for output file (optional)
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx not installed. Run: pip install python-docx"
            )

        self.output_path = Path(output_path) if output_path else None
        self._doc = Document()
        self._setup_rtl_support()

        app_logger.info("WordExporter initialized")

    def _setup_rtl_support(self) -> None:
        """Setup RTL support for Arabic text."""
        # Set default paragraph direction to RTL
        try:
            styles = self._doc.styles
            style = styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(12)

            # Set RTL for paragraph
            pPr = style.element.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            pPr.append(bidi)

        except Exception as e:
            app_logger.warning(f"RTL setup warning: {e}")

    def _set_cell_rtl(self, cell) -> None:
        """Set RTL direction for table cell."""
        try:
            tcPr = cell._tc.get_or_add_tcPr()
            # Set text direction RTL
            textDirection = OxmlElement('w:textDirection')
            textDirection.set(qn('w:val'), 'rlTb')  # right-to-left, top-to-bottom
            tcPr.append(textDirection)
        except Exception:
            pass

    def add_heading(
        self,
        text: str,
        level: int = 1,
        alignment: str = "right"
    ) -> None:
        """
        Add heading to document.

        Args:
            text: Heading text
            level: Heading level (0-9, 0=Title)
            alignment: "left", "center", "right"
        """
        heading = self._doc.add_heading(text, level=level)

        # Set alignment
        if alignment == "center":
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "left":
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:  # right (default for Arabic)
            heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def add_paragraph(
        self,
        text: str,
        bold: bool = False,
        italic: bool = False,
        alignment: str = "right",
        font_size: int = 12
    ) -> None:
        """
        Add paragraph to document.

        Args:
            text: Paragraph text
            bold: Bold text
            italic: Italic text
            alignment: "left", "center", "right", "justify"
            font_size: Font size in points
        """
        para = self._doc.add_paragraph()

        # Set alignment
        alignments = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        para.alignment = alignments.get(alignment, WD_ALIGN_PARAGRAPH.RIGHT)

        # Add run with formatting
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(font_size)

    def add_table(
        self,
        data: List[Dict[str, Any]],
        headers: Optional[List[str]] = None,
        style: str = "Table Grid"
    ) -> None:
        """
        Add table to document.

        Args:
            data: List of row dictionaries
            headers: Column headers (uses dict keys if not provided)
            style: Table style name
        """
        if not data:
            return

        # Get headers from first row if not provided
        if headers is None:
            headers = list(data[0].keys())

        # Create table
        table = self._doc.add_table(rows=1, cols=len(headers))
        table.style = style
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = str(header)
            self._set_cell_rtl(header_cells[i])
            # Bold header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add data rows
        for row_data in data:
            row_cells = table.add_row().cells
            for i, header in enumerate(headers):
                value = row_data.get(header, "")
                row_cells[i].text = str(value) if value is not None else ""
                self._set_cell_rtl(row_cells[i])

    def add_line_break(self, count: int = 1) -> None:
        """Add empty paragraphs."""
        for _ in range(count):
            self._doc.add_paragraph()

    def add_page_break(self) -> None:
        """Add page break."""
        self._doc.add_page_break()

    def add_horizontal_line(self) -> None:
        """Add horizontal line."""
        para = self._doc.add_paragraph()
        para.add_run("─" * 50)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_image(
        self,
        image_path: str,
        width: Optional[float] = None,
        height: Optional[float] = None
    ) -> bool:
        """
        Add image to document.

        Args:
            image_path: Path to image file
            width: Width in inches (optional)
            height: Height in inches (optional)

        Returns:
            True if successful
        """
        try:
            if width:
                self._doc.add_picture(image_path, width=Inches(width))
            elif height:
                self._doc.add_picture(image_path, height=Inches(height))
            else:
                self._doc.add_picture(image_path)
            return True
        except Exception as e:
            app_logger.error(f"Failed to add image: {e}")
            return False

    def add_header(self, text: str) -> None:
        """Add document header."""
        section = self._doc.sections[0]
        header = section.header
        para = header.paragraphs[0]
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_footer(self, text: str) -> None:
        """Add document footer."""
        section = self._doc.sections[0]
        footer = section.footer
        para = footer.paragraphs[0]
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def save(self, output_path: Optional[str] = None) -> bool:
        """
        Save document to file.

        Args:
            output_path: Output path (uses init path if not provided)

        Returns:
            True if successful
        """
        try:
            path = Path(output_path) if output_path else self.output_path

            if path is None:
                app_logger.error("No output path specified")
                return False

            # Ensure directory exists
            path.parent.mkdir(parents=True, exist_ok=True)

            self._doc.save(str(path))
            app_logger.info(f"Document saved: {path}")
            return True

        except Exception as e:
            app_logger.error(f"Failed to save document: {e}")
            return False

    def get_document(self) -> Document:
        """Get the underlying Document object."""
        return self._doc


def create_employee_report(
    employee: Dict[str, Any],
    output_path: str,
    include_photo: bool = False
) -> bool:
    """
    Create employee report document.

    Args:
        employee: Employee data dictionary
        output_path: Output file path
        include_photo: Include employee photo

    Returns:
        True if successful
    """
    try:
        doc = WordExporter(output_path)

        # Title
        doc.add_heading("بطاقة معلومات الموظف", level=0)
        doc.add_line_break()

        # Date
        doc.add_paragraph(
            f"تاريخ التقرير: {datetime.now().strftime('%Y-%m-%d')}",
            alignment="left"
        )
        doc.add_horizontal_line()
        doc.add_line_break()

        # Employee info as table
        info_data = []

        field_labels = {
            "name_ar": "الاسم",
            "name_en": "Name",
            "employee_number": "رقم الموظف",
            "department": "القسم",
            "job_title": "المسمى الوظيفي",
            "hire_date": "تاريخ التعيين",
            "salary": "الراتب",
            "phone": "الهاتف",
            "email": "البريد الإلكتروني",
            "nationality": "الجنسية",
            "iban": "رقم الحساب البنكي"
        }

        for key, label in field_labels.items():
            if key in employee and employee[key]:
                info_data.append({
                    "البيان": label,
                    "القيمة": str(employee[key])
                })

        doc.add_table(info_data, headers=["البيان", "القيمة"])

        # Footer
        doc.add_line_break(2)
        doc.add_horizontal_line()
        doc.add_paragraph(
            "INTEGRA - نظام الإدارة المتكامل",
            alignment="center",
            font_size=10
        )

        return doc.save()

    except Exception as e:
        app_logger.error(f"Failed to create employee report: {e}")
        return False


def create_employees_list_report(
    employees: List[Dict[str, Any]],
    output_path: str,
    title: str = "قائمة الموظفين",
    columns: Optional[List[str]] = None
) -> bool:
    """
    Create employees list report.

    Args:
        employees: List of employee dictionaries
        output_path: Output file path
        title: Report title
        columns: Columns to include (optional)

    Returns:
        True if successful
    """
    try:
        doc = WordExporter(output_path)

        # Title
        doc.add_heading(title, level=0)
        doc.add_paragraph(
            f"تاريخ التقرير: {datetime.now().strftime('%Y-%m-%d')}",
            alignment="left"
        )
        doc.add_paragraph(
            f"عدد الموظفين: {len(employees)}",
            alignment="left"
        )
        doc.add_line_break()

        # Default columns
        if columns is None:
            columns = ["الاسم", "القسم", "المسمى الوظيفي", "الراتب"]

        # Filter data to include only specified columns
        filtered_data = []
        for emp in employees:
            row = {col: emp.get(col, "") for col in columns}
            filtered_data.append(row)

        doc.add_table(filtered_data, headers=columns)

        # Footer
        doc.add_line_break(2)
        doc.add_paragraph(
            "INTEGRA - نظام الإدارة المتكامل",
            alignment="center",
            font_size=10
        )

        return doc.save()

    except Exception as e:
        app_logger.error(f"Failed to create employees list report: {e}")
        return False
